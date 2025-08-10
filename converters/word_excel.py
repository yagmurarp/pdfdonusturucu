from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.utils import get_column_letter

def docx_to_xlsx(src_docx, dst_xlsx):
    doc = Document(src_docx)

    # Collect tables; also capture paragraphs
    tables = doc.tables
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    wb = Workbook()
    ws0 = wb.active
    ws0.title = "Paragraflar"
    if paragraphs:
        ws0.append(["Paragraflar"])
        for p in paragraphs:
            ws0.append([p])
    else:
        ws0.append(["Paragraf bulunamadı veya boş."])

    # Each docx table -> one sheet
    for ti, table in enumerate(tables, start=1):
        ws = wb.create_sheet(title=f"Tablo{ti}")
        # Determine max columns by first row
        max_cols = 0
        for row in table.rows:
            max_cols = max(max_cols, len(row.cells))

        # Write header-like first row? We'll just dump raw cells
        for r_idx, row in enumerate(table.rows, start=1):
            for c_idx, cell in enumerate(row.cells, start=1):
                ws.cell(row=r_idx, column=c_idx, value=cell.text)

        # Optional: set column widths
        for c in range(1, max_cols+1):
            ws.column_dimensions[get_column_letter(c)].width = 20

    wb.save(dst_xlsx)

def xlsx_to_docx(src_xlsx, dst_docx):
    wb = load_workbook(src_xlsx)
    doc = Document()

    for si, sheet in enumerate(wb.worksheets, start=1):
        if si > 1:
            doc.add_page_break()
        doc.add_heading(f"Sayfa: {sheet.title}", level=1)

        # Find used range
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            doc.add_paragraph("(Boş sayfa)")
            continue

        # Create a table with the size of used range
        max_cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = "Table Grid"

        for r_idx, row in enumerate(rows):
            for c_idx in range(max_cols):
                val = row[c_idx] if c_idx < len(row) else ""
                table.cell(r_idx, c_idx).text = "" if val is None else str(val)

    doc.save(dst_docx)